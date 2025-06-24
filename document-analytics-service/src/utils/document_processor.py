import os
import re
import PyPDF2
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.pipeline import Pipeline
import nltk
from datetime import datetime

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX processing will be limited.")

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

class DocumentProcessor:
    def __init__(self):
        self.classifier = None
        self.categories = ['Academic', 'Business', 'Technical', 'Legal', 'Medical', 'General']
        self._initialize_classifier()
    
    def extract_text_from_pdf(self, file_path):
        """Extract text content from PDF file - Enhanced version"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                total_pages = len(pdf_reader.pages)

                print(f"Extracting text from PDF: {file_path}, Total pages: {total_pages}")

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        print(f"Page {page_num + 1}: extracted {len(page_text)} characters")
                    else:
                        print(f"Page {page_num + 1}: no text extracted")

                final_text = text.strip()
                print(f"Total text extracted: {len(final_text)} characters")
                return final_text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path):
        """Extract text content from DOCX file - Enhanced version"""
        try:
            if DOCX_AVAILABLE:
                doc = DocxDocument(file_path)
                text = ""
                paragraph_count = 0

                print(f"Extracting text from DOCX: {file_path}")

                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Only add non-empty paragraphs
                        text += paragraph.text + "\n"
                        paragraph_count += 1

                # Also extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + " "
                    text += "\n"

                final_text = text.strip()
                print(f"DOCX extraction complete: {paragraph_count} paragraphs, {len(final_text)} characters")
                return final_text
            else:
                # Fallback when python-docx is not available
                print(f"python-docx not available for file: {file_path}")
                return f"DOCX file: {os.path.basename(file_path)}"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return f"DOCX file: {os.path.basename(file_path)}"
    
    def extract_title_from_pdf(self, file_path):
        """Extract title from PDF metadata or content"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Try to get title from metadata
                if pdf_reader.metadata and pdf_reader.metadata.title:
                    return pdf_reader.metadata.title.strip()
                
                # Try to extract title from first page content
                if len(pdf_reader.pages) > 0:
                    first_page_text = pdf_reader.pages[0].extract_text()
                    lines = first_page_text.split('\n')
                    for line in lines[:10]:  # Check first 10 lines
                        line = line.strip()
                        if len(line) > 10 and len(line) < 200:  # Reasonable title length
                            return line
                
                # Fallback to filename
                return os.path.splitext(os.path.basename(file_path))[0]
        except Exception as e:
            print(f"Error extracting title from PDF: {e}")
            return os.path.splitext(os.path.basename(file_path))[0]
    
    def extract_title_from_docx(self, file_path):
        """Extract title from DOCX file"""
        try:
            if DOCX_AVAILABLE:
                doc = DocxDocument(file_path)

                # Try to get title from document properties
                if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                    return doc.core_properties.title.strip()

                # Try to extract title from first paragraph
                if doc.paragraphs:
                    for paragraph in doc.paragraphs[:5]:  # Check first 5 paragraphs
                        text = paragraph.text.strip()
                        if len(text) > 10 and len(text) < 200:  # Reasonable title length
                            return text

            # Fallback to filename
            return os.path.splitext(os.path.basename(file_path))[0]
        except Exception as e:
            print(f"Error extracting title from DOCX: {e}")
            return os.path.splitext(os.path.basename(file_path))[0]
    
    def extract_metadata_from_pdf(self, file_path):
        """Extract metadata from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {}
                
                if pdf_reader.metadata:
                    metadata['author'] = pdf_reader.metadata.author if pdf_reader.metadata.author else None
                    metadata['creation_date'] = pdf_reader.metadata.creation_date if pdf_reader.metadata.creation_date else None
                    metadata['last_modified'] = pdf_reader.metadata.modification_date if pdf_reader.metadata.modification_date else None
                
                return metadata
        except Exception as e:
            print(f"Error extracting metadata from PDF: {e}")
            return {}
    
    def extract_metadata_from_docx(self, file_path):
        """Extract metadata from DOCX file - simplified version"""
        try:
            # Return basic metadata for DOCX files
            stat = os.stat(file_path)
            metadata = {
                'author': None,
                'creation_date': datetime.fromtimestamp(getattr(stat, 'st_birthtime', stat.st_ctime)),
                'last_modified': datetime.fromtimestamp(stat.st_mtime)
            }
            return metadata
        except Exception as e:
            print(f"Error extracting metadata from DOCX: {e}")
            return {}
    
    def _initialize_classifier(self):
        """Initialize the document classifier with sample training data"""
        # Sample training data for different categories
        training_data = [
            ("research methodology analysis statistical significant", "Academic"),
            ("university college student education learning", "Academic"),
            ("business strategy market revenue profit", "Business"),
            ("company management financial report quarterly", "Business"),
            ("algorithm software programming code development", "Technical"),
            ("system architecture database network security", "Technical"),
            ("contract agreement legal terms conditions", "Legal"),
            ("court case law regulation compliance", "Legal"),
            ("medical patient treatment diagnosis therapy", "Medical"),
            ("health clinical study pharmaceutical drug", "Medical"),
            ("general information document text content", "General"),
            ("various topics discussion overview summary", "General")
        ]
        
        texts = [item[0] for item in training_data]
        labels = [item[1] for item in training_data]
        
        # Create and train the classifier
        self.classifier = Pipeline([
            ('tfidf', TfidfVectorizer(stop_words='english', max_features=1000)),
            ('classifier', MultinomialNB())
        ])
        
        self.classifier.fit(texts, labels)
    
    def classify_document(self, text):
        """Classify document based on its content"""
        if not text or not self.classifier:
            return "General", 0.5
        
        try:
            # Clean and preprocess text
            cleaned_text = re.sub(r'[^a-zA-Z\s]', '', text.lower())
            
            # Predict category and confidence
            prediction = self.classifier.predict([cleaned_text])[0]
            probabilities = self.classifier.predict_proba([cleaned_text])[0]
            confidence = max(probabilities)
            
            return prediction, confidence
        except Exception as e:
            print(f"Error classifying document: {e}")
            return "General", 0.5
    
    def highlight_text(self, text, search_terms):
        """Highlight search terms (words, phrases, or sentences) in text"""
        if not search_terms or not text:
            return text

        highlighted_text = text
        print(f"Highlighting text with terms: {search_terms}")  # Debug

        # Sort search terms by length (longest first) to avoid partial highlighting
        sorted_terms = sorted(search_terms, key=len, reverse=True)

        for term in sorted_terms:
            if not term.strip():
                continue

            # Escape special regex characters
            escaped_term = re.escape(term.strip())
            print(f"Highlighting term: '{term}' (escaped: '{escaped_term}')")  # Debug

            def replace_func(match):
                highlighted = f"<mark>{match.group()}</mark>"
                print(f"Replacing '{match.group()}' with '{highlighted}'")  # Debug
                return highlighted

            # Apply highlighting with case-insensitive matching
            before_highlight = highlighted_text
            highlighted_text = re.sub(escaped_term, replace_func, highlighted_text, flags=re.IGNORECASE)

            # Check if highlighting actually happened
            if before_highlight != highlighted_text:
                print(f"Successfully highlighted '{term}' in text")
            else:
                print(f"No matches found for '{term}' in text")

        print(f"Final highlighted text length: {len(highlighted_text)}")  # Debug
        return highlighted_text
    
    def extract_match_contexts(self, text, search_terms, lines_before=1, lines_after=1):
        """Extract context around matches (3 lines total: 1 before, match line, 1 after)"""
        if not text or not search_terms:
            return []

        # Split text into lines
        lines = text.split('\n')
        contexts = []

        for term in search_terms:
            term_lower = term.lower()

            # Find all lines containing the term
            for i, line in enumerate(lines):
                if term_lower in line.lower():
                    # Calculate context range
                    start_idx = max(0, i - lines_before)
                    end_idx = min(len(lines), i + lines_after + 1)

                    # Extract context lines
                    context_lines = lines[start_idx:end_idx]
                    context_text = '\n'.join(context_lines)

                    # Highlight the term in context
                    highlighted_context = self.highlight_text(context_text, [term])

                    contexts.append({
                        'term': term,
                        'line_number': i + 1,
                        'context': highlighted_context,
                        'context_start_line': start_idx + 1,
                        'context_end_line': end_idx,
                        'match_line_in_context': i - start_idx
                    })

        return contexts

    def search_documents(self, documents, keywords):
        """Search documents for keywords, phrases, and sentences - Enhanced version with contexts"""
        if not keywords:
            return documents

        matching_docs = []
        search_query = keywords.strip()
        search_query_lower = search_query.lower()

        # Split into individual words for fallback search
        individual_words = [word.strip() for word in keywords.split() if word.strip()]

        print(f"Searching for: '{search_query}' in {len(documents)} documents")  # Debug

        for doc in documents:
            content_text = doc.get('content_text', '')
            title = doc.get('title', '')

            # Debug: Print content length
            print(f"Document '{title}' content length: {len(content_text)} characters")

            # Search in both content and title (case-insensitive)
            full_content = (content_text + ' ' + title).lower()

            found_matches = []
            match_type = None

            # First: Try to find the exact phrase/sentence (case-insensitive)
            if search_query_lower in full_content:
                found_matches.append(search_query)  # Keep original case for highlighting
                match_type = 'exact_phrase'
                print(f"Found exact phrase '{search_query}' in document '{title}'")
            else:
                # Second: If no exact phrase match, search for individual words
                for word in individual_words:
                    word_lower = word.lower()
                    if word_lower in full_content:
                        found_matches.append(word)  # Keep original case for highlighting
                        print(f"Found word '{word}' in document '{title}'")
                if found_matches:
                    match_type = 'individual_words'

            # If any matches found, add to results
            if found_matches:
                # For highlighting, use the original search query if it's an exact match
                highlight_terms = [search_query] if match_type == 'exact_phrase' else found_matches

                # Extract contexts around matches
                match_contexts = self.extract_match_contexts(content_text, highlight_terms)

                highlighted_content = self.highlight_text(content_text, highlight_terms)
                highlighted_title = self.highlight_text(title, highlight_terms)

                doc_copy = doc.copy()
                doc_copy['highlighted_content'] = highlighted_content
                doc_copy['highlighted_title'] = highlighted_title
                doc_copy['matched_terms'] = found_matches
                doc_copy['match_type'] = match_type
                doc_copy['search_query'] = search_query
                doc_copy['match_contexts'] = match_contexts  # Add contexts
                doc_copy['total_matches'] = len(match_contexts)  # Add total count
                doc_copy['content_preview'] = content_text[:500] + '...' if len(content_text) > 500 else content_text
                matching_docs.append(doc_copy)

        print(f"Search completed. Found {len(matching_docs)} matching documents")  # Debug
        return matching_docs

